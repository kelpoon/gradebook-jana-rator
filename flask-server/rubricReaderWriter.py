from __future__ import print_function
import sys
import subprocess

subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'pydrive','-q','--no-color'])
subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'python-docx','-q','--no-color'])
subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'docx','-q','--no-color'])
subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'termcolor','-q','--no-color'])
subprocess.check_call([sys.executable, '-m', 'pip', 'install', '--upgrade','google-api-python-client',
                       ' google-auth-httplib2','google-auth-oauthlib','-q','--no-color'])

from pydrive.auth import GoogleAuth
from pydrive.drive import GoogleDrive
import os
import os.path
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
from google.oauth2 import service_account
from docx.api import Document
import docx
from docx.enum.text import WD_COLOR_INDEX
from termcolor import colored
from pathlib import Path

# *************************************************************** #
# Input variables

lightColor = 7 #yellow means partial credit
darkColor = 4 #green means full credit

# Default folder and spreadsheet IDs
# folderLink = '1tyeoX2ZgmgSW57cS8z-r4S_Z6hKPl2yo'
# spreadsheetLink = '1-CZJljjGtjP9J2AX5n80jlOdPWBCvum5pm7F8rDS510'


folderRoot = 'drive_download'

# *************************************************************** #


# *************************************************************** #
# Functions #

# Drive related

def authenticate():
    """
    Authenticates the loads google drive for folder input
    :params: none
    :return: authorized Google Drive
    """
    gauth = GoogleAuth()

    # Tries to load saved client credentials
    gauth.LoadCredentialsFile("credentials.txt")

    if gauth.credentials is None:
        # Authenticate if they're not there
        gauth.LocalWebserverAuth()
    elif gauth.access_token_expired:
        # Refresh them if expired
        gauth.Refresh()
    else:
        # Initialize the saved credentials
        gauth.Authorize()

    # Save the current credentials to a file
    gauth.SaveCredentialsFile("credentials.txt")

    drive = GoogleDrive(gauth)
    return drive

def escape_fname(name):
    """
    Simple function to reformat folder names
    :params: foldername
    :return: reformatted foldername
    """
    return name.replace('/','_')

def search_folder(folder_id, root):
    """
    Iterates through google drive folder until it finds google doc files to download
    If it finds relevant files, it downloads them to your local machine
    :param folder_id: URL ID of google drive folder
    :param root: computer root for download, i.e where the files should download
    :return: none, but downloads files
    """
    drive = authenticate()
    f = open("failed.txt","w+")

    file_list = drive.ListFile({'q': "'%s' in parents and trashed=false" % folder_id}).GetList()
    for file in file_list:

        if file['mimeType'].split('.')[-1] == 'folder':
            foldername = escape_fname(file['title'])
            create_folder(root,foldername)
            search_folder(file['id'], '{}{}/'.format(root,foldername))
        else:
            download_mimetype = None
            filename = escape_fname(file['title'])
            filename = '{}{}'.format(root,filename)
            try:
                print('DOWNLOADING:', filename)
                if file['mimeType'] in MIMETYPES:
                    download_mimetype = MIMETYPES[file['mimeType']]

                    file.GetContentFile(filename+EXTENSIONS[file['mimeType']], mimetype=download_mimetype)
                else:
                    file.GetContentFile(filename)
            except:
                print('FAILED')
                f.write(filename+'\n')

def create_folder(path,name):
    """
    Creates local file to house downloaded templates
    :param name: name of created folder
    :param path: computer root for download, i.e where the files should download
    :return: none
    """
    os.mkdir('{}{}'.format(path,escape_fname(name)))


# Terminal related

# for outputting pretty colors to the terminal
def convert_wd_color_index_to_termcolor(color_index):
    """
    Converts color ID to string colors if recognized, else returns white
    :param name: name of created folder
    :param path: computer root for download, i.e where the files should download
    :return: none
    """
    if (color_index == WD_COLOR_INDEX.BLUE):
        return "blue"
    if (color_index == WD_COLOR_INDEX.TURQUOISE):
        return "cyan"
    if (color_index == WD_COLOR_INDEX.GREEN):
        return "green"
    if (color_index == WD_COLOR_INDEX.BRIGHT_GREEN):
        return "light_green"
    if (color_index == WD_COLOR_INDEX.RED):
        return "red"
    if (color_index == WD_COLOR_INDEX.YELLOW):
        return "yellow"

    print("WARNING: Unrecognized color index: %s" % (color_index))
    return "white"


def calculateScoreFromHighlights(highlights):
    """
    Calculates total numerical score by summing highlights
    :param highlights: array of all highlights and corresponding scores
    :return: numerical score
    """
    score = 0
    for h in highlights:
        score += h[1]
    return score

# *************************************************************** #

SHOW_EVERYTHING_INCLUDING_NON_HIGHLIGHTED = False

# START OF DRIVE THINGS #

MIMETYPES = {
        # Drive Document files as MS dox
        'application/vnd.google-apps.document': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        # Drive Sheets files as MS Excel files.
        'application/vnd.google-apps.spreadsheet': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        # Drive presentation as MS pptx
        'application/vnd.google-apps.presentation': 'application/vnd.openxmlformats-officedocument.presentationml.presentation'
        # see https://developers.google.com/drive/v3/web/mime-types
    }
EXTENSIONS = {
        'application/vnd.google-apps.document': '.docx',
        'application/vnd.google-apps.spreadsheet': '.xlsx',
        'application/vnd.google-apps.presentation': '.pptx'
}

def run_rubricReaderWriter(folder_url, gradebook_url):
# creating the folder, downloading files from drive
    drive = authenticate()

    f = open("failed.txt","w+")
    folder_id = folder_url
    root = folderRoot
    if not os.path.exists(root):
        os.makedirs(root)

    search_folder(folder_id,root+'/')
    f.close() 


    # END OF DRIVE THINGS #

    dir_list = os.listdir(root)
    print(dir_list)


    # If modifying these scopes, delete the file token.json.


    SCOPES = ['https://www.googleapis.com/auth/spreadsheets']
    SERVICE_ACCOUNT_FILE = 'keys.json'
    creds = service_account.Credentials.from_service_account_file(
        SERVICE_ACCOUNT_FILE, scopes=SCOPES)
    service = build('sheets', 'v4', credentials=creds)
    sheet = service.spreadsheets()

    # Start of doc things

    # The big for loop which iterates through all the downloaded files in the folder. Count is used to help iterate through some 
    # of the arrays.

    count = 0

    for file in Path(root).iterdir():
        name = os.listdir(root)[count][:-46] #46 charas is the "end year linear algebra rubric"

        cell_range = "GRADES!B" + str(count+6) #specific to the gradebook Jana uses (grades tab and B6 onward)

        document = docx.Document(file)

        # For debugging

        # print(colored("\n========= Found %d tables in the document ==========" % (len(document.tables)), "blue"))

        # # Print out summary of tables found in the document
        # for t, table in enumerate(document.tables):
        #     print(colored("Table %d has %d rows and %d columns" % (t, len(table.rows), len(table.columns)), "yellow"))



        # each of these lists will contain tuples of (text, score) which we'll later remove dupes using set
        allFoundational = []
        allProficients = []
        allExemplarys = []


        # Print out detailed contents of each table, along with what is highlighted
        for t, table in enumerate(document.tables):
            # lists for counting within each table (has dupes)
            highlightedFoundationals = []
            highlightedProficients = []
            highlightedExemplarys = []

            # print("\n\nTABLE %d:" % (t))
            for r, row in enumerate(table.rows): #goes through rows in each table. NOTE: something is broken with this, number is off so we end up w/ dupes
                # print("\n-------- Row %d --------" % (r))
                for c, cell in enumerate(row.cells): 
                    # print()
                    for p, paragraph in enumerate(cell.paragraphs): #each (content) standard in a cell is a paragraph
                        if (len(paragraph.runs) == 0): #skip cell if no paragraphs
                            continue
                        # score will be the precentage of runs inside this paragraph that are highlighted
                        text = paragraph.text
                        for r2, run in enumerate(paragraph.runs): #using runs to determine highlighted colors within the paragraphs. We found that when something is highlighted multiple colors, it will split into multiple runs
                            colors_foundational = []
                            colors_proficient = []
                            colors_exemplary = []
                            if (c == 1): #1st column is foundational
                                for i in range(len(paragraph.runs)): #still within single paragraph. just getting all the different colors
                                    colors_foundational.append(paragraph.runs[i].font.highlight_color)  
                                                    
                            if (c == 2): #2nd proficient
                                for i in range(len(paragraph.runs)):
                                    colors_proficient.append(paragraph.runs[i].font.highlight_color)

                            if (c == 3): #3rd exemplary
                                for i in range(len(paragraph.runs)):
                                    colors_exemplary.append(paragraph.runs[i].font.highlight_color)
                            #set data type gets rid of all duplicates
                            colors_foundational = list(set(colors_foundational))
                            colors_proficient = list(set(colors_proficient))
                            colors_exemplary = list(set(colors_exemplary))
                            
                            #manually checking cases to determine score for the paragraph
                            if len(colors_foundational) == 1:
                                print(run.text)
                                print(colors_foundational)
                                if colors_foundational[0] == darkColor:
                                    highlightedFoundationals.append((text,1))
                                if colors_foundational[0] == lightColor:
                                    highlightedFoundationals.append((text,.5))
                                
                            elif len(colors_foundational) > 1:
                                print(run.text)
                                print(colors_foundational)
                                if darkColor in colors_foundational and lightColor in colors_foundational:
                                    highlightedFoundationals.append((text,.75))
                                elif darkColor in colors_foundational:
                                    highlightedFoundationals.append((text,.5))
                                elif lightColor in colors_foundational:
                                    highlightedFoundationals.append((text,.25))
                            

                            if len(colors_proficient) == 1:
                                if colors_proficient[0] == darkColor:
                                    highlightedProficients.append((text,1))
                                if colors_proficient[0] == lightColor:
                                    highlightedProficients.append((text,.5))
                                
                            elif len(colors_proficient) > 1:
                                if darkColor in colors_proficient and lightColor in colors_proficient:
                                    highlightedProficients.append((text,.75))
                                elif darkColor in colors_proficient:
                                    highlightedProficients.append((text,.5))
                                elif lightColor in colors_proficient:
                                    highlightedProficients.append((text,.25))

                            if len(colors_exemplary) == 1:
                                if colors_exemplary[0] == darkColor:
                                    highlightedExemplarys.append((text,1))
                                if colors_exemplary[0] == lightColor:
                                    highlightedExemplarys.append((text,.5))

                                
                            elif len(colors_exemplary) > 1:
                                if darkColor in colors_exemplary and lightColor in colors_exemplary:
                                    highlightedExemplarys.append((text,.75))
                                elif darkColor in colors_exemplary:
                                    highlightedExemplarys.append((text,.5))
                                elif lightColor in colors_exemplary:
                                    highlightedExemplarys.append((text,.25))

                        # More debug code
                        # for r2, run in enumerate(paragraph.runs):
                        #     if run.font.highlight_color is not None:
                        #         print(
                        #             colored("*Table %d, Row %d, Cell %d, Paragraph %d, Run %d: %s" % (t, r, c, p, r2, run.text),
                        #                     convert_wd_color_index_to_termcolor(run.font.highlight_color)))
                        #     else:
                        #         if SHOW_EVERYTHING_INCLUDING_NON_HIGHLIGHTED:
                        #             print(" Table %d, Row %d, Cell %d, Paragraph %d, Run %d: %s" % (t, r, c, p, r2, run.text))

            # put all the scores in one list
            allFoundational.append(highlightedFoundationals)
            allProficients.append(highlightedProficients)
            allExemplarys.append(highlightedExemplarys)

        # ****
        # Console stuff for testing
        # ***

        # print("\n\n========= FOUNDATIONALS ==========")
        # # highlightedFoundationals = list(set(highlightedFoundationals))
        # # print(highlightedFoundationals)
        # # print(calculateScoreFromHighlights(highlightedFoundationals))
        # for i in allFoundational:
        #     print(calculateScoreFromHighlights(list(set(i))))


        # print("\n\n========= PROFICIENTS ==========")
        # # highlightedProficients = list(set(highlightedProficients))
        # # # print(highlightedProficients)
        # # print(calculateScoreFromHighlights(highlightedProficients))
        # for i in allProficients:
        #     print(calculateScoreFromHighlights(list(set(i))))

        # print("\n\n========= EXEMPLARYS ==========")
        # # highlightedExemplarys = list(set(highlightedExemplarys))
        # # # print(highlightedExemplarys)
        # # print(calculateScoreFromHighlights(highlightedExemplarys))
        # for i in allExemplarys:
        #     print(calculateScoreFromHighlights(list(set(i))))

        #getting rid of duplicates again and calculating scores for content, habits, skills etc.
        content_f = calculateScoreFromHighlights(list(set(allFoundational[0])))
        content_p = calculateScoreFromHighlights(list(set(allProficients[0])))
        content_e = calculateScoreFromHighlights(list(set(allExemplarys[0])))
        skills_f = calculateScoreFromHighlights(list(set(allFoundational[1])))
        skills_p = calculateScoreFromHighlights(list(set(allProficients[1])))
        skills_e = calculateScoreFromHighlights(list(set(allExemplarys[1])))
        habits_f = calculateScoreFromHighlights(list(set(allFoundational[2])))
        habits_p = calculateScoreFromHighlights(list(set(allProficients[2])))
        habits_e = calculateScoreFromHighlights(list(set(allExemplarys[2])))

        score = [[name,None,None,content_f,content_p,content_e,None,None,skills_f,skills_p,skills_e,None,None,habits_f,habits_p,habits_e]]
        aoa = [["1/1/2020",4000]]
        #updating the google sheet
        request = sheet.values().update(spreadsheetId=gradebook_url,
                                        range=cell_range,valueInputOption="USER_ENTERED",
                                        body = {"values":score}).execute()
        count +=1
